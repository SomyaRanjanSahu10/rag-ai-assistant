"""Document text extraction — PDF, DOCX, TXT, images with OCR."""

import os
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text_from_file(file_path: str) -> tuple[str, dict]:
    path = Path(file_path)
    ext  = path.suffix.lower()
    meta: dict = {"filename": path.name, "file_size": os.path.getsize(file_path), "file_type": ext}

    if ext == ".pdf":
        text, m = _pdf(file_path)
    elif ext == ".docx":
        text, m = _docx(file_path)
    elif ext == ".txt":
        text = _txt(file_path)
        m    = {}
    elif ext in {".png", ".jpg", ".jpeg", ".webp", ".bmp"}:
        text, m = _image(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    meta.update(m)
    meta["extracted_length"] = len(text)
    if not text.strip():
        raise ValueError(f"No text extracted from {path.name}")
    return text, meta


# ── PDF ────────────────────────────────────────────────────────────────────────

def _pdf(path: str) -> tuple[str, dict]:
    import PyPDF2
    meta: dict = {"method": "pypdf2"}
    parts: list[str] = []

    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        meta["page_count"] = len(reader.pages)
        if reader.metadata:
            meta["title"]  = reader.metadata.title or ""
            meta["author"] = reader.metadata.author or ""
        for i, page in enumerate(reader.pages):
            t = page.extract_text() or ""
            if t.strip():
                parts.append(f"[Page {i+1}]\n{t}")

    text = "\n\n".join(parts)

    # Fallback to OCR if text is sparse
    if len(text) < meta["page_count"] * 80:
        try:
            ocr_text = _pdf_ocr(path)
            if len(ocr_text) > len(text):
                meta["method"] = "ocr_fallback"
                return ocr_text, meta
        except Exception as exc:
            logger.warning("PDF OCR fallback failed: %s", exc)

    return text, meta


def _pdf_ocr(path: str) -> str:
    try:
        import pdf2image, pytesseract
        images = pdf2image.convert_from_path(path, dpi=200)
        return "\n\n".join(
            f"[Page {i+1}]\n{pytesseract.image_to_string(img, lang='eng')}"
            for i, img in enumerate(images)
        )
    except Exception:
        import pdf2image, easyocr, numpy as np
        reader = easyocr.Reader(["en"], gpu=False)
        images = pdf2image.convert_from_path(path, dpi=150)
        return "\n\n".join(
            f"[Page {i+1}]\n{' '.join(reader.readtext(np.array(img), detail=0))}"
            for i, img in enumerate(images)
        )


# ── DOCX ───────────────────────────────────────────────────────────────────────

def _docx(path: str) -> tuple[str, dict]:
    from docx import Document
    doc   = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            prefix = "### " if para.style.name.startswith("Heading") else ""
            parts.append(prefix + para.text)
    for table in doc.tables:
        rows = [" | ".join(c.text.strip() for c in r.cells) for r in table.rows]
        parts.append("[Table]\n" + "\n".join(r for r in rows if r.strip()))
    return "\n".join(parts), {"method": "python-docx", "paragraph_count": len(doc.paragraphs)}


# ── TXT ────────────────────────────────────────────────────────────────────────

def _txt(path: str) -> str:
    for enc in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
        try:
            with open(path, encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    with open(path, "rb") as f:
        return f.read().decode("utf-8", errors="replace")


# ── Images ─────────────────────────────────────────────────────────────────────

def _image(path: str) -> tuple[str, dict]:
    text = ""
    method = "none"

    try:
        import pytesseract
        from PIL import Image
        img = Image.open(path).convert("RGB")
        text = pytesseract.image_to_string(img, config="--oem 3 --psm 3", lang="eng")
        method = "tesseract"
    except Exception as exc:
        logger.warning("Tesseract failed: %s", exc)

    if len(text.strip()) < 50:
        try:
            import easyocr, numpy as np
            from PIL import Image
            reader = easyocr.Reader(["en"], gpu=False)
            arr    = np.array(Image.open(path))
            items  = reader.readtext(arr)
            ocr_t  = "\n".join(item[1] for item in items if item[2] > 0.4)
            if len(ocr_t) > len(text):
                text   = ocr_t
                method = "easyocr"
        except Exception as exc:
            logger.warning("EasyOCR failed: %s", exc)

    if not text.strip():
        raise ValueError("No text could be extracted from image")

    return text, {"method": method}
